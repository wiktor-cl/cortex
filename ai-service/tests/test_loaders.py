from io import BytesIO

import pytest
from docx import Document as DocxDocument

from app.ingestion import loaders
from app.ingestion.loaders import UnsupportedFileTypeError, load_document


class _FakePage:
    def __init__(self, text: str) -> None:
        self._text = text

    def extract_text(self) -> str:
        return self._text


class _FakeReader:
    def __init__(self, data: bytes) -> None:
        self.pages = [_FakePage("Page one content."), _FakePage(""), _FakePage("Page three content.")]


def test_load_pdf_reports_one_based_page_numbers_and_skips_blank_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(loaders, "PdfReader", _FakeReader)

    pages = load_document("manual.pdf", b"irrelevant")

    # The blank second page is dropped, but the surviving pages keep their true 1-based page number.
    assert [(p.text, p.page_number) for p in pages] == [
        ("Page one content.", 1),
        ("Page three content.", 3),
    ]


def test_load_txt_returns_single_page_without_page_number() -> None:
    pages = load_document("manual.txt", b"Hello world.")

    assert len(pages) == 1
    assert pages[0].text == "Hello world."
    assert pages[0].page_number is None


def test_load_markdown_strips_formatting_to_plain_text() -> None:
    pages = load_document("manual.md", b"# Title\n\nSome **bold** text.")

    assert len(pages) == 1
    assert "Title" in pages[0].text
    assert "**" not in pages[0].text


def test_load_html_extracts_text_content() -> None:
    pages = load_document("manual.html", b"<html><body><p>Torque: 12 Nm</p></body></html>")

    assert len(pages) == 1
    assert "Torque: 12 Nm" in pages[0].text


def test_load_docx_extracts_paragraph_text() -> None:
    document = DocxDocument()
    document.add_paragraph("Safety procedure step one.")
    document.add_paragraph("Safety procedure step two.")
    buffer = BytesIO()
    document.save(buffer)

    pages = load_document("manual.docx", buffer.getvalue())

    assert len(pages) == 1
    assert "Safety procedure step one." in pages[0].text
    assert "Safety procedure step two." in pages[0].text


def test_load_document_unknown_extension_raises() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        load_document("manual.exe", b"binary")


def test_load_document_no_extension_raises() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        load_document("manual", b"data")
